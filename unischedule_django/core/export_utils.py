from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.styles import ParagraphStyle
from django.http import FileResponse
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
HOURS = ["08:00","09:00","10:00","11:00","12:00","13:00","14:00","15:00","16:00","17:00"]

from reportlab.lib.styles import ParagraphStyle

def generate_pdf(schedule, failed, semester, faculty, session, institution):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter),
                            leftMargin=30, rightMargin=30, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    
    # Use smaller body style for better wrapping
    small_style = ParagraphStyle(
        name='SmallCell',
        fontSize=5.5,
        leading=6,
        spaceAfter=0,
    )

    story = []

    header_style = ParagraphStyle(
    name="Header",
    fontSize=14,
    leading=18,
    alignment=TA_CENTER,
    spaceAfter=6,
    )
    print(institution)
    institution_name = institution.upper()  # Or pass this dynamically if needed

    story.append(Paragraph(f"<b>{institution_name}</b>", header_style))
    story.append(Paragraph(f"{session} Academic Session", header_style))
    story.append(Paragraph(f"Faculty of {faculty} – General Timetable", header_style))
    story.append(Spacer(1, 12))
    # Table header
    table_data = [["Days"] + HOURS]

    # Table body rows
    for day in DAYS:
        row = [Paragraph(day, small_style)]
        for hour in HOURS:
            booked = [
                Paragraph(
                    f"<b>{b['course_id']}</b><br/>{b['lecturer']}<br/><i>Room: {b['room']}</i>\n",
                    small_style
                )
                for b in schedule if b['day'] == day and hour >= b['start_time'] and hour < b['end_time']
            ]
            row.append(booked if booked else "")
        table_data.append(row)

    # Adjust column widths to squeeze 11 columns to one page
    col_widths = [0.9 * inch] + [0.82 * inch] * len(HOURS)

    tbl = Table(table_data, repeatRows=1, colWidths=col_widths)
    tbl.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkgray),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 2),
        ('RIGHTPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))
    story.append(tbl)

    if failed:
        story.append(Spacer(1, 12))
        story.append(Paragraph("Failed Scheduling Attempts:", styles["Heading2"]))
        for f in failed:
            story.append(Paragraph(f"❌ {f}", styles["Normal"]))

    doc.build(story)
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename="Optimized_Schedule.pdf")

def generate_docx(schedule, failed, semester, year, dept, faculty, session):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(8)

    # Add title
    title = f"{dept}, {faculty} - {semester} ({year})"
    subtitle = f"{session}"
    document.add_heading(title, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph(subtitle).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph()

    # Create table
    num_cols = len(HOURS) + 1
    table = document.add_table(rows=len(DAYS) + 1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Set uniform width
    col_width = Inches(0.9)  # fit all columns nicely
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    # Header row
    table.cell(0, 0).text = "Days"
    for i, hour in enumerate(HOURS):
        table.cell(0, i + 1).text = hour

    # Table content
    for row_idx, day in enumerate(DAYS):
        table.cell(row_idx + 1, 0).text = day
        for col_idx, hour in enumerate(HOURS):
            cell = table.cell(row_idx + 1, col_idx + 1)
            entries = [
                f"{b['course_id']} - {b['course_name']}\n{b['lecturer']}\nRoom: {b['room']}"
                for b in schedule if b['day'] == day and hour >= b['start_time'] and hour < b['end_time']
            ]
            cell.text = "\n\n".join(entries) if entries else ""

    # Format cells: shrink padding & wrap text
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = paragraph.runs
                if run:
                    run[0].font.size = Pt(8)

    # Failed bookings
    if failed:
        document.add_page_break()
        document.add_heading("Failed Scheduling Attempts:", level=2)
        for f in failed:
            document.add_paragraph(f"❌ {f}")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename="Optimized_Schedule.docx")
